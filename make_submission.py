import shutil, copy
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx

SRC = '/root/.claude/uploads/e267424e-8088-5218-af77-098b94456e93/c8679a50-SEBANG_AI_______________.docx'
OUT = '/home/user/dongbu-contract-performance/SEBANG_AI_공모전_신청서_작성본.docx'

doc = Document(SRC)

# ── 헬퍼 ──────────────────────────────────────────────────────────
def set_cell(cell, text, bold=False, size=None, align=None, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_run(para, text, bold=False, size=None, color=None):
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def table_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '999999')
        tblBorders.append(el)
    tblPr.append(tblBorders)

NAVY   = (0x0F, 0x29, 0x5C)
WHITE  = (0xFF, 0xFF, 0xFF)
GRAY_BG = 'F2F2F2'
NAVY_HEX = '0F295C'

# ── 새 문서 ────────────────────────────────────────────────────────
d = Document()
d.core_properties.author = ''

# 기본 여백
for sec in d.sections:
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

styles = d.styles
normal = styles['Normal']
normal.font.name = '맑은 고딕'
normal.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════
# 첨부#1 제목
# ══════════════════════════════════════════════════════════════════
h = d.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h.add_run('[첨부#1] 신청서')
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(*NAVY)

t = d.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t.add_run('SEBANG AI 아이디어 공모전 신청서')
r2.bold = True; r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(*NAVY)
d.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# 1. 신청자 정보
# ══════════════════════════════════════════════════════════════════
s1 = d.add_paragraph()
add_run(s1, '1. 신청자 정보', bold=True, size=11, color=NAVY)

tbl1 = d.add_table(rows=8, cols=4)
tbl1.style = 'Table Grid'
table_border(tbl1)

def mrow(tbl, row_idx, label, val, label_bg=GRAY_BG):
    row = tbl.rows[row_idx]
    set_cell(row.cells[0], label, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row.cells[0], label_bg)
    # merge cols 1-3
    merged = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
    set_cell(merged, val, size=9)

# 참가 구분 (span 전체)
r0 = tbl1.rows[0]
set_cell(r0.cells[0], '참가 구분', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(r0.cells[0], GRAY_BG)
m0 = r0.cells[1].merge(r0.cells[2]).merge(r0.cells[3])
set_cell(m0, '☑ 개인    □ 단체(팀)  총 인원 (         )명', size=9)

# 팀 대표/개인 헤더 (2행: span rows for "팀 대표 및 개인" label)
r1 = tbl1.rows[1]
set_cell(r1.cells[0], '팀 대표\n및 개인', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(r1.cells[0], GRAY_BG)
set_cell(r1.cells[1], '소속', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(r1.cells[1], GRAY_BG)
mc1 = r1.cells[2].merge(r1.cells[3])
set_cell(mc1, '', size=9)

r2 = tbl1.rows[2]
# merge row cells[0] with above? — keep separate, just blank
set_cell(r2.cells[0], '', size=9)
set_cell_bg(r2.cells[0], GRAY_BG)
set_cell(r2.cells[1], '이름', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(r2.cells[1], GRAY_BG)
mc2 = r2.cells[2].merge(r2.cells[3])
set_cell(mc2, '', size=9)

r3 = tbl1.rows[3]
set_cell(r3.cells[0], '', size=9)
set_cell_bg(r3.cells[0], GRAY_BG)
set_cell(r3.cells[1], '이메일', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(r3.cells[1], GRAY_BG)
set_cell(r3.cells[2], '직급', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(r3.cells[2], GRAY_BG)
set_cell(r3.cells[3], '', size=9)

r4 = tbl1.rows[4]
set_cell(r4.cells[0], '', size=9)
set_cell_bg(r4.cells[0], GRAY_BG)
set_cell(r4.cells[1], '연락처', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(r4.cells[1], GRAY_BG)
mc4 = r4.cells[2].merge(r4.cells[3])
set_cell(mc4, '', size=9)

# 팀원 (3행)
for idx, row_i in enumerate([5,6,7]):
    rr = tbl1.rows[row_i]
    label = '팀원' if idx==0 else ''
    set_cell(rr.cells[0], label, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(rr.cells[0], GRAY_BG)
    set_cell(rr.cells[1], '소속', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(rr.cells[1], GRAY_BG)
    set_cell(rr.cells[2], '이름', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(rr.cells[2], GRAY_BG)
    set_cell(rr.cells[3], '', size=9)

d.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# 2. 제안 개요
# ══════════════════════════════════════════════════════════════════
s2 = d.add_paragraph()
add_run(s2, '2. 제안 개요', bold=True, size=11, color=NAVY)

tbl2 = d.add_table(rows=3, cols=2)
tbl2.style = 'Table Grid'
table_border(tbl2)
tbl2.columns[0].width = Cm(3)

# 제목
ra = tbl2.rows[0]
set_cell(ra.cells[0], '제목', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(ra.cells[0], GRAY_BG)
set_cell(ra.cells[1],
    '화물차 안전운임 즉시 조회 웹/모바일 서비스\n'
    '— 구간별·차종별 법정 최저운임 원스톱 검색',
    bold=True, size=10)

# 분야
rb = tbl2.rows[1]
set_cell(rb.cells[0], '분야', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(rb.cells[0], GRAY_BG)
rb.cells[1].text = ''
p_b = rb.cells[1].paragraphs[0]
add_run(p_b, '□ 업무 혁신 : 반복되거나 수작업 비중이 높은 비효율적인 업무 자동화 (효율화)\n', size=9)
add_run(p_b, '□ 의사결정 지원 : 보고 자료 등 작성 시, 신속하고 정확한 의사결정을 지원하는 분석 툴\n', size=9)
add_run(p_b, '☑ 기타 : 고객 서비스 개선 및 신규 가치 창출 등 기타 아이디어 제안', size=9, bold=True)

# 제안 개요
rc = tbl2.rows[2]
set_cell(rc.cells[0], '제안\n개요', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(rc.cells[0], GRAY_BG)
rc.cells[1].text = ''
p_c = rc.cells[1].paragraphs[0]
add_run(p_c,
    '화물차 안전운임제는 구간·차종·품목별로 운임 기준이 복잡하고 국토교통부 고시가 수시로 개정되어, '
    '현장 담당자가 정확한 운임을 즉시 확인하기 어려운 상황이다.\n'
    '출발지·도착지·차종을 입력하면 해당 구간의 법정 최저운임을 실시간으로 조회·비교할 수 있는 웹/모바일 서비스를 개발하였다.\n'
    '자동완성 검색, 모바일 최적화 UI, 최신 고시 데이터 반영 구조를 적용하여 조회 시간 단축 및 운임 계산 오류를 최소화하였다.\n'
    '사내 운영팀, 협력 운수사, 화주 담당자 모두 언제 어디서든 스마트폰으로 정확한 운임 기준을 확인할 수 있어 업무 효율과 대외 신뢰도 향상이 기대된다.',
    size=9)

d.add_paragraph()

# 유의사항
note = d.add_paragraph()
note.paragraph_format.left_indent = Cm(0.5)
add_run(note, '※ 유의사항', bold=True, size=8, color=(0x60,0x60,0x60))
notes = [
    '1) 제안서 제출 시 ① 신청서, ② 아이디어 활용 동의서를 모두 제출해야 하며, 제출 서류는 반환하지 않습니다.',
    '2) 신청자는 공모안이 제3자의 저작권, 초상권 등을 침해하지 않도록 주의 의무를 준수해야 하며, 관련된 사항은 신청자의 책임으로 합니다.',
    '3) 타 공모전 입상·수상작 또는 표절·중복응모·도용 등 부정행위 발생 시 심사 제외, 수상 취소 및 향후 당사에서 진행하는 DT 및 AI 관련 모든 공모 응시 참여가 제한됩니다.',
    '4) 접수 현황 및 심사 결과에 따라 시상 규모 축소 또는 변경할 수 있습니다.',
    '5) 1차 심사 및 평가는 신청자 정보를 미공개로 진행할 예정입니다.',
]
for n in notes:
    np = d.add_paragraph()
    np.paragraph_format.left_indent = Cm(0.8)
    add_run(np, n, size=8, color=(0x60,0x60,0x60))

d.add_paragraph()
agree = d.add_paragraph()
agree.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(agree,
    '신청자 본인은 위 내용에 동의하며, < SEBANG AI 아이디어 공모전 >에 참가 신청합니다.',
    bold=True, size=10)
d.add_paragraph()

date_p = d.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(date_p, '2026년         월         일', size=11)
d.add_paragraph()

sign_tbl = d.add_table(rows=1, cols=3)
sign_tbl.alignment = 1  # center
table_border(sign_tbl)
set_cell(sign_tbl.rows[0].cells[0], '신청인(팀 대표)', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(sign_tbl.rows[0].cells[0], GRAY_BG)
set_cell(sign_tbl.rows[0].cells[1], '(소속)\n\n', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell(sign_tbl.rows[0].cells[2], '(이름)                  (인)', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

d.add_paragraph()
d.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3. 아이디어 제안서
# ══════════════════════════════════════════════════════════════════
s3 = d.add_paragraph()
add_run(s3, '3. 아이디어 제안서', bold=True, size=11, color=NAVY)

idea_tbl = d.add_table(rows=5, cols=2)
idea_tbl.style = 'Table Grid'
table_border(idea_tbl)
idea_tbl.columns[0].width = Cm(2.8)

def idea_row(tbl, row_i, label, content_fn):
    row = tbl.rows[row_i]
    set_cell(row.cells[0], label, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row.cells[0], GRAY_BG)
    row.cells[1].text = ''
    content_fn(row.cells[1])

# ① 제목
def fill_title(cell):
    p = cell.paragraphs[0]
    add_run(p, '화물차 안전운임 즉시 조회 웹/모바일 서비스\n— 구간별·차종별 법정 최저운임 원스톱 검색', bold=True, size=10)

# ② 현황 진단
def fill_current(cell):
    p = cell.paragraphs[0]
    add_run(p, '① 업무 현황 및 문제점\n', bold=True, size=9, color=NAVY)
    add_run(p,
        '• 화물차 안전운임제(컨테이너·시멘트 품목 대상)는 출발지·도착지 구간, 차종(톤급), 품목에 따라 운임 기준이 세분화되어 있으며, 국토교통부 고시가 연 1회 이상 개정된다.\n'
        '• 현재 운임 기준 확인은 ①국토교통부 고시 PDF를 직접 검색하거나 ②담당자가 관리하는 엑셀 파일을 공유받는 방식에 의존하고 있어, 조회 1건당 평균 5~10분 이상 소요된다.\n'
        '• 엑셀 파일은 버전 관리가 되지 않아 구버전 운임 기준으로 계약·청구가 이루어지는 오류가 반복적으로 발생한다.\n'
        '• 현장 운전기사, 협력 운수사, 화주 담당자, 사내 운영팀 등 다수의 이해관계자가 각기 다른 자료를 참고하여 분쟁 및 재작업이 발생한다.\n'
        '• 모바일 환경에서 즉시 조회 가능한 공식 도구가 없어 현장 대응력이 저하된다.',
        size=9)

# ③ 제안 내용
def fill_proposal(cell):
    p = cell.paragraphs[0]
    add_run(p, '① 서비스 개요\n', bold=True, size=9, color=NAVY)
    add_run(p,
        '출발지·도착지·차종(톤급)·품목을 선택하면 해당 구간의 화물차 안전운임(법정 최저운임)을 즉시 조회할 수 있는 웹/모바일 서비스이다.\n\n',
        size=9)
    add_run(p, '② 주요 기능\n', bold=True, size=9, color=NAVY)
    features = [
        ('스마트 검색 (자동완성)', '출발지·도착지 입력 시 자동완성 드롭다운 제공, 키보드(↑↓ Enter) 및 터치 모두 지원'),
        ('구간별 운임 조회', '출발지·도착지·차종·품목 조합으로 현행 법정 최저운임 즉시 표시'),
        ('데이터 최신화 구조', '국토교통부 고시 개정 시 관리자 1회 업데이트만으로 전체 사용자에게 즉시 반영'),
        ('모바일 최적화', '반응형 레이아웃으로 스마트폰·태블릿에서도 편리하게 사용'),
        ('운임 이력 비교', '고시 변경 전후 운임 이력을 함께 제공하여 계약·협상 시 참고 가능'),
    ]
    for fn, fd in features:
        add_run(p, f'  • {fn} : ', bold=True, size=9)
        add_run(p, f'{fd}\n', size=9)
    add_run(p, '\n③ 장애요인 및 해결방안\n', bold=True, size=9, color=NAVY)
    issues = [
        ('고시 데이터 구조화 작업 필요', '기존 PDF·엑셀 데이터를 1회 DB화 후 고시 개정 시 부분 수정만으로 유지'),
        ('구간 수가 많아 검색 불편', '자동완성 + 부분 문자열 매칭 검색으로 빠른 선택 지원'),
        ('모바일 환경 접근성', 'PWA 방식 적용으로 앱 설치 없이 모바일 홈화면 추가 및 오프라인 캐싱 가능'),
    ]
    for iss, sol in issues:
        add_run(p, f'  • {iss} → ', bold=True, size=9)
        add_run(p, f'{sol}\n', size=9)
    add_run(p, '\n④ 서비스 구성도 (개념)\n', bold=True, size=9, color=NAVY)
    add_run(p,
        '  [사용자 : 운전기사 / 운영팀 / 화주]\n'
        '          ↓ 출발지·도착지·차종 입력\n'
        '  [웹/모바일 프론트엔드]\n'
        '    - 자동완성 검색 UI  /  운임 결과 카드 표시\n'
        '          ↓ API 조회\n'
        '  [안전운임 데이터베이스]\n'
        '    - 고시 기준 구간별 운임 테이블  /  개정 이력 관리\n'
        '          ↑ 관리자 업데이트\n'
        '  [국토교통부 안전운임 고시 데이터]',
        size=9)

# ④ 활용 도구
def fill_tools(cell):
    p = cell.paragraphs[0]
    tools = [
        ('AI 활용', 'Claude AI — 서비스 설계, 코드 생성, 데이터 구조화 보조'),
        ('개발 도구', 'HTML5 / CSS3 / JavaScript (단일 파일 웹앱), Supabase (클라우드 DB·API)'),
        ('데이터', '국토교통부 화물자동차 안전운임 고시 공공 데이터'),
        ('배포', 'GitHub Pages (별도 서버 비용 없이 즉시 배포·운영 가능)'),
        ('접근성', 'PWA 지원 — 앱 설치 없이 모바일 홈화면 아이콘 추가, 오프라인 캐싱'),
    ]
    for k, v in tools:
        add_run(p, f'• {k} : ', bold=True, size=9)
        add_run(p, f'{v}\n', size=9)

# ⑤ 기대 효과
def fill_effect(cell):
    p = cell.paragraphs[0]
    effects = [
        ('업무 효율', '운임 조회 소요시간 5~10분 → 30초 이내 단축 (약 90% 절감)'),
        ('오류 감소', '구버전 운임 기준 적용으로 인한 청구·계약 오류 제거'),
        ('현장 대응력', '현장 기사·운영팀·화주가 스마트폰으로 즉시 운임 확인, 즉석 협의 가능'),
        ('비용 절감', '별도 서버 없이 운영 가능(GitHub Pages + Supabase 무료 티어)하여 구축·유지 비용 최소화'),
        ('확장성', '향후 자동 적정운임 계산기, 운임 이상 감지, 청구서 자동 생성 등으로 발전 가능'),
    ]
    for k, v in effects:
        add_run(p, f'• {k} : ', bold=True, size=9)
        add_run(p, f'{v}\n', size=9)

idea_row(idea_tbl, 0, '제목', fill_title)
idea_row(idea_tbl, 1, '현황\n진단', fill_current)
idea_row(idea_tbl, 2, '제안\n내용', fill_proposal)
idea_row(idea_tbl, 3, '활용\n도구', fill_tools)
idea_row(idea_tbl, 4, '기대\n효과', fill_effect)

d.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 첨부#2 아이디어 활용 동의서
# ══════════════════════════════════════════════════════════════════
h2 = d.add_paragraph()
h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(h2, '[첨부#2] 아이디어 활용 동의서', bold=True, size=13, color=NAVY)
d.add_paragraph()

t2 = d.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(t2, '아이디어 활용 동의서', bold=True, size=14, color=NAVY)
d.add_paragraph()

intro = d.add_paragraph()
add_run(intro,
    '본인은 세방㈜에서 주관하는 <SEBANG AI 아이디어 공모전>(이하 "공모전")에 아이디어를 제출함에 있어, 아래 사항에 동의합니다.',
    size=10)
d.add_paragraph()

below = d.add_paragraph()
below.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(below, '-     아          래     -', bold=True, size=10)
d.add_paragraph()

consents = [
    '본인이 제출한 아이디어는 회사의 업무 개선, AI 활용 및 업무 자동화 과제 추진, 교육, 홍보 및 기타 내부 목적을 위해 활용될 수 있다.',
    '회사는 해당 아이디어를 사용, 수정, 보완, 재구성할 수 있으며, 아이디어를 기반으로 한 결과물을 내부자료(보고서, 사례집 등)로 제작 및 활용할 수 있다.',
    '회사는 제출된 아이디어를 시스템 개선, AI 기반 서비스 기획 및 개발 등 다양한 형태로 제한 없이 활용할 수 있다.',
]
for c in consents:
    cp = d.add_paragraph(style='List Bullet')
    cp.paragraph_format.left_indent = Cm(1)
    add_run(cp, c, size=10)

d.add_paragraph()
add_run(d.add_paragraph(), '본인은 상기 내용을 충분히 이해하였으며, 이에 동의합니다.', bold=True, size=10)
d.add_paragraph()

date_p2 = d.add_paragraph()
date_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(date_p2, '2026년         월         일', size=11)
d.add_paragraph()

sign_tbl2 = d.add_table(rows=4, cols=3)
sign_tbl2.alignment = 1
table_border(sign_tbl2)
roles = ['대표자', '팀원', '팀원', '팀원']
for i, role in enumerate(roles):
    row = sign_tbl2.rows[i]
    set_cell(row.cells[0], role, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row.cells[0], GRAY_BG)
    set_cell(row.cells[1], '(소속)\n\n', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[2], '(이름)                  (인)', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

d.save(OUT)
print('저장 완료:', OUT)
